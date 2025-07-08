# app.py
from flask import Flask, render_template, jsonify, request, send_from_directory, redirect, url_for, session
import os
import sqlite3
import hashlib
from datetime import datetime
from flask_cors import CORS
import jdatetime
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.parse
import shutil
from functools import wraps
import collections

# --- Configuration ---
TENANT_DB_BASE_DIR = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'data', 'tenants')
GENERATED_LETTERS_BASE_DIR = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'generated_letters')
COMPANY_TEMPLATES_BASE_DIR = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'company_templates')


# Initialize the Flask application
app = Flask(__name__)
CORS(app)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'a_very_secret_key_for_development_and_sessions')

# --- Superadmin Credentials (Hardcoded) ---
SUPERADMIN_USERNAME = "superadmin"
SUPERADMIN_PASSWORD = "your_strong_password" 

# Ensure necessary directories exist at startup
os.makedirs(TENANT_DB_BASE_DIR, exist_ok=True)
os.makedirs(GENERATED_LETTERS_BASE_DIR, exist_ok=True)
os.makedirs(COMPANY_TEMPLATES_BASE_DIR, exist_ok=True)

# --- Decorator for Superadmin Authentication ---
def superadmin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'superadmin_logged_in' not in session:
            return redirect(url_for('serve_superadmin_login_page'))
        return f(*args, **kwargs)
    return decorated_function

def api_superadmin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'superadmin_logged_in' not in session:
            return jsonify({"message": "Unauthorized access"}), 401
        return f(*args, **kwargs)
    return decorated_function

# --- Database Functions ---
def get_db_path(company_name):
    """Constructs the database path for a given company."""
    clean_company_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
    if not clean_company_name:
        raise ValueError("Company name cannot be empty or contain only special characters.")
    db_dir = os.path.join(TENANT_DB_BASE_DIR, clean_company_name)
    os.makedirs(db_dir, exist_ok=True)
    return os.path.join(db_dir, f'{clean_company_name}.db')

def init_db(company_name):
    """Initializes the database schema for a new company."""
    db_path = get_db_path(company_name)
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'user'
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS organizations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            industry TEXT,
            phone TEXT,
            email TEXT,
            address TEXT,
            description TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS contacts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            organization_id INTEGER,
            first_name TEXT NOT NULL,
            last_name TEXT NOT NULL,
            title TEXT,
            phone TEXT,
            email TEXT,
            notes TEXT,
            FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS letters (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_abbr TEXT NOT NULL,
            seq_num INTEGER NOT NULL,
            letter_code_persian TEXT NOT NULL UNIQUE,
            type TEXT NOT NULL,
            date_shamsi_persian TEXT NOT NULL,
            subject TEXT NOT NULL,
            body TEXT NOT NULL,
            organization_id INTEGER,
            contact_id INTEGER,
            local_file_path TEXT,
            current_gregorian_date TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            FOREIGN KEY (organization_id) REFERENCES organizations(id) ON DELETE SET NULL,
            FOREIGN KEY (contact_id) REFERENCES contacts(id) ON DELETE SET NULL,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS company_settings (
            company_name TEXT PRIMARY KEY,
            company_short_name TEXT,
            company_full_name_footer TEXT,
            letter_template_path TEXT
        )
    ''')
    conn.commit()
    conn.close()

def get_db_connection(company_name):
    """Establishes a database connection for a given company."""
    db_path = get_db_path(company_name)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

# --- Helper function for placeholder replacement ---
def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    """Replaces a placeholder in a paragraph, even if it spans multiple runs."""
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        replacement_text = str(replacement) if replacement is not None else ""
        new_text = full_text.replace(placeholder, replacement_text)
        paragraph.clear()
        paragraph.add_run(new_text)
        return True
    return False

# --- Routes for serving HTML files ---
@app.route('/')
def serve_login_page():
    """Serves the login.html page."""
    return render_template('login.html')

@app.route('/main_app')
def serve_main_app():
    """Serves the index.html (main application) page."""
    return render_template('index.html')

# --- Superadmin Routes ---
@app.route('/superadmin/login')
def serve_superadmin_login_page():
    """Serves the superadmin login page."""
    return render_template('superadmin_login.html')

@app.route('/superadmin')
@superadmin_required
def serve_superadmin_panel():
    """Serves the protected superadmin.html panel."""
    return render_template('superadmin.html')

@app.route('/superadmin/logout')
def superadmin_logout():
    """Logs out the superadmin by clearing the session."""
    session.pop('superadmin_logged_in', None)
    return redirect(url_for('serve_superadmin_login_page'))


# --- Superadmin API Endpoints ---
@app.route('/api/superadmin/login', methods=['POST'])
def superadmin_login():
    """Handles superadmin login."""
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    if username == SUPERADMIN_USERNAME and password == SUPERADMIN_PASSWORD:
        session['superadmin_logged_in'] = True
        return jsonify({"message": "Login successful"}), 200
    else:
        return jsonify({"message": "Invalid credentials"}), 401

@app.route('/api/superadmin/create_company', methods=['POST'])
@api_superadmin_required
def create_company():
    """Superadmin endpoint to create a new company and its admin user."""
    data = request.get_json()
    company_name = data.get('company_name')
    admin_email = data.get('admin_email')
    admin_password = data.get('admin_password')

    if not all([company_name, admin_email, admin_password]):
        return jsonify({"message": "Company name, admin email, and password are required"}), 400

    try:
        if os.path.exists(get_db_path(company_name)):
            return jsonify({"message": f"Company '{company_name}' already exists"}), 409
    except ValueError as ve:
        return jsonify({"message": str(ve)}), 400

    try:
        init_db(company_name)
        conn = sqlite3.connect(get_db_path(company_name))
        cursor = conn.cursor()
        password_hash = hashlib.sha256(admin_password.encode()).hexdigest()
        cursor.execute("INSERT INTO users (email, password_hash, role) VALUES (?, ?, ?)",
                       (admin_email, password_hash, 'admin'))
        
        cursor.execute("INSERT INTO company_settings (company_name, company_short_name, company_full_name_footer, letter_template_path) VALUES (?, ?, ?, ?)",
                       (company_name, company_name, f"شرکت {company_name}", None))
        
        conn.commit()
        conn.close()
        return jsonify({"message": f"Company '{company_name}' created with admin user '{admin_email}'"}), 201
    except sqlite3.IntegrityError:
        return jsonify({"message": "Admin user with this email already exists for this company"}), 409
    except Exception as e:
        print(f"Error creating company: {e}")
        return jsonify({"message": f"Error creating company: {str(e)}"}), 500

@app.route('/api/superadmin/companies', methods=['GET'])
@api_superadmin_required
def get_all_companies():
    """Superadmin endpoint to get a list of all existing companies."""
    try:
        if not os.path.exists(TENANT_DB_BASE_DIR):
            return jsonify([]), 200
        
        company_dirs = [d for d in os.listdir(TENANT_DB_BASE_DIR) if os.path.isdir(os.path.join(TENANT_DB_BASE_DIR, d))]
        companies = [{"name": name} for name in company_dirs]
        return jsonify(companies), 200
    except Exception as e:
        print(f"Error fetching companies list: {e}")
        return jsonify({"message": f"Error fetching companies list: {str(e)}"}), 500

@app.route('/api/superadmin/companies/<string:company_name>', methods=['DELETE'])
@api_superadmin_required
def delete_company_data(company_name):
    """Superadmin endpoint to delete a company and all its associated data."""
    if not company_name:
        return jsonify({"message": "Company name is required"}), 400

    clean_company_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
    if not clean_company_name or clean_company_name != company_name:
        return jsonify({"message": "Invalid company name format"}), 400

    try:
        db_dir_path = os.path.join(TENANT_DB_BASE_DIR, clean_company_name)
        letters_dir_path = os.path.join(GENERATED_LETTERS_BASE_DIR, clean_company_name)
        templates_dir_path = os.path.join(COMPANY_TEMPLATES_BASE_DIR, clean_company_name)

        if os.path.exists(db_dir_path):
            shutil.rmtree(db_dir_path)
        if os.path.exists(letters_dir_path):
            shutil.rmtree(letters_dir_path)
        if os.path.exists(templates_dir_path):
            shutil.rmtree(templates_dir_path)

        return jsonify({"message": f"Company '{clean_company_name}' and all its data have been deleted successfully."}), 200
    except Exception as e:
        error_message = f"An error occurred while deleting company '{clean_company_name}': {str(e)}"
        print(error_message)
        return jsonify({"message": error_message}), 500


# --- User/Authentication Endpoints ---
@app.route('/api/login', methods=['POST'])
def login():
    """Handles user login."""
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    company_name = data.get('company_name')

    if not all([email, password, company_name]):
        return jsonify({"message": "Email, password, and company name are required"}), 400

    try:
        db_path = get_db_path(company_name)
        if not os.path.exists(db_path):
            return jsonify({"message": f"Company '{company_name}' not found"}), 404
    except ValueError as ve:
        return jsonify({"message": str(ve)}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("SELECT id, email, password_hash, role FROM users WHERE email = ?", (email,))
        user = cursor.fetchone()
        conn.close()

        if user and hashlib.sha256(password.encode()).hexdigest() == user['password_hash']:
            conn_settings = get_db_connection(company_name)
            cursor_settings = conn_settings.cursor()
            cursor_settings.execute("SELECT company_short_name, company_full_name_footer, letter_template_path FROM company_settings WHERE company_name = ?", (company_name,))
            settings = cursor_settings.fetchone()
            conn_settings.close()

            user_data = {
                "user_id": user['id'],
                "user_email": user['email'],
                "role": user['role'],
                "company_name": company_name,
                "company_short_name": settings['company_short_name'] if settings else company_name,
                "company_full_name_footer": settings['company_full_name_footer'] if settings else f"شرکت {company_name}",
                "letter_template_path": settings['letter_template_path'] if settings else None
            }
            return jsonify(user_data), 200
        else:
            return jsonify({"message": "Invalid credentials"}), 401
    except Exception as e:
        print(f"Error during login: {e}")
        return jsonify({"message": f"Error during login: {str(e)}"}), 500

@app.route('/api/user/change-password', methods=['POST'])
def change_password():
    """Allows a logged-in user to change their password."""
    data = request.get_json()
    user_id = data.get('user_id')
    company_name = data.get('company_name')
    current_password = data.get('current_password')
    new_password = data.get('new_password')

    if not all([user_id, company_name, current_password, new_password]):
        return jsonify({"message": "All fields are required"}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE id = ?", (user_id,))
        user = cursor.fetchone()

        if not user:
            conn.close()
            return jsonify({"message": "User not found"}), 404

        if hashlib.sha256(current_password.encode()).hexdigest() != user['password_hash']:
            conn.close()
            return jsonify({"message": "رمز عبور فعلی اشتباه است"}), 403

        new_password_hash = hashlib.sha256(new_password.encode()).hexdigest()
        cursor.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_password_hash, user_id))
        conn.commit()
        conn.close()

        return jsonify({"message": "رمز عبور با موفقیت تغییر کرد"}), 200

    except Exception as e:
        print(f"Error changing password: {e}")
        return jsonify({"message": f"Error changing password: {str(e)}"}), 500

# --- User Management (Admin only) ---
@app.route('/api/users', methods=['POST'])
def add_user():
    """Adds a new user to a company."""
    data = request.get_json()
    company_name = data.get('company_name')
    email = data.get('email')
    password = data.get('password')
    role = data.get('role', 'user')

    if not all([company_name, email, password]):
        return jsonify({"message": "Company name, email, and password are required"}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        password_hash = hashlib.sha256(password.encode()).hexdigest()
        cursor.execute("INSERT INTO users (email, password_hash, role) VALUES (?, ?, ?)",
                       (email, password_hash, role))
        conn.commit()
        new_user_id = cursor.lastrowid
        conn.close()
        return jsonify({"message": "User added successfully", "user_id": new_user_id, "email": email, "role": role}), 201
    except sqlite3.IntegrityError:
        return jsonify({"message": "User with this email already exists in this company"}), 409
    except Exception as e:
        print(f"Error adding user: {e}")
        return jsonify({"message": f"Error adding user: {str(e)}"}), 500

@app.route('/api/users', methods=['GET'])
def get_users():
    """Retrieves all users for a company."""
    company_name = request.args.get('company_name')
    if not company_name:
        return jsonify({"message": "Company name is required"}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("SELECT id, email, role FROM users")
        users = cursor.fetchall()
        conn.close()
        return jsonify([dict(user) for user in users]), 200
    except Exception as e:
        print(f"Error fetching users: {e}")
        return jsonify({"message": f"Error fetching users: {str(e)}"}), 500

@app.route('/api/users/<int:user_id>', methods=['GET', 'PUT', 'DELETE'])
def handle_user(user_id):
    """Handles GET, PUT, and DELETE requests for a single user."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        if not company_name:
            return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("SELECT id, email, role FROM users WHERE id = ?", (user_id,))
            user = cursor.fetchone()
            conn.close()
            if user:
                return jsonify(dict(user)), 200
            else:
                return jsonify({"message": "User not found"}), 404
        except Exception as e:
            return jsonify({"message": f"Error fetching user: {str(e)}"}), 500

    elif request.method == 'PUT':
        data = request.get_json()
        company_name = data.get('company_name')
        new_role = data.get('role')
        if not all([company_name, new_role]):
            return jsonify({"message": "Company name and new role are required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("UPDATE users SET role = ? WHERE id = ?", (new_role, user_id))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "User not found"}), 404
            
            cursor.execute("SELECT id, email, role FROM users WHERE id = ?", (user_id,))
            updated_user = cursor.fetchone()
            conn.close()
            return jsonify(dict(updated_user)), 200
        except Exception as e:
            return jsonify({"message": f"Error updating user role: {str(e)}"}), 500

    elif request.method == 'DELETE':
        company_name = request.args.get('company_name')
        if not company_name:
            return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("DELETE FROM users WHERE id = ?", (user_id,))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "User not found"}), 404
            conn.close()
            return jsonify({"message": "User deleted successfully"}), 200
        except Exception as e:
            return jsonify({"message": f"Error deleting user: {str(e)}"}), 500


# --- Settings Endpoints ---
@app.route('/api/settings', methods=['GET', 'POST'])
def handle_settings():
    """Handles GET and POST requests for company settings."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        if not company_name:
            return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("SELECT company_short_name, company_full_name_footer, letter_template_path FROM company_settings WHERE company_name = ?", (company_name,))
            settings = cursor.fetchone()
            conn.close()
            if settings:
                return jsonify(dict(settings)), 200
            else:
                return jsonify({
                    "company_short_name": company_name,
                    "company_full_name_footer": f"شرکت {company_name}",
                    "letter_template_path": None
                }), 200
        except Exception as e:
            return jsonify({"message": f"Error fetching company settings: {str(e)}"}), 500

    if request.method == 'POST':
        data = request.get_json()
        company_name = data.get('company_name')
        company_short_name = data.get('company_short_name')
        company_full_name_footer = data.get('company_full_name_footer')
        if not company_name:
            return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("UPDATE company_settings SET company_short_name = ?, company_full_name_footer = ? WHERE company_name = ?",
                           (company_short_name, company_full_name_footer, company_name))
            if cursor.rowcount == 0:
                cursor.execute("INSERT INTO company_settings (company_name, company_short_name, company_full_name_footer) VALUES (?, ?, ?)",
                               (company_name, company_short_name, company_full_name_footer))
            conn.commit()
            conn.close()
            return jsonify({"message": "Company settings updated successfully"}), 200
        except Exception as e:
            return jsonify({"message": f"Error updating company settings: {str(e)}"}), 500

@app.route('/api/settings/upload_template', methods=['POST'])
def upload_letter_template():
    """Handles uploading a DOCX letter template for a company."""
    company_name = request.form.get('company_name')
    if 'letter_template' not in request.files:
        return jsonify({"message": "No letter_template file part"}), 400
    file = request.files['letter_template']
    if file.filename == '':
        return jsonify({"message": "No selected file"}), 400
    if not company_name:
        return jsonify({"message": "Company name is required"}), 400
    if not file.filename.lower().endswith('.docx'):
        return jsonify({"message": "Invalid file type. Only .docx files are allowed."}), 400
    
    company_template_dir = os.path.join(COMPANY_TEMPLATES_BASE_DIR, company_name)
    os.makedirs(company_template_dir, exist_ok=True)
    filename = os.path.basename(file.filename)
    template_path = os.path.join(company_template_dir, filename)
    try:
        file.save(template_path)
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("UPDATE company_settings SET letter_template_path = ? WHERE company_name = ?",
                       (template_path, company_name))
        if cursor.rowcount == 0:
            cursor.execute("INSERT INTO company_settings (company_name, letter_template_path) VALUES (?, ?)",
                           (company_name, template_path))
        conn.commit()
        conn.close()
        return jsonify({
            "message": "Letter template uploaded and saved successfully",
            "template_filename": filename,
            "template_full_path": template_path
        }), 200
    except Exception as e:
        return jsonify({"message": f"Error uploading letter template: {str(e)}"}), 500


# --- Organization Endpoints ---
@app.route('/api/organizations', methods=['GET', 'POST'])
def handle_organizations():
    """Handles GET and POST requests for organizations."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        search_term = request.args.get('search', '')
        if not company_name:
            return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            query = "SELECT * FROM organizations"
            params = []
            if search_term:
                query += " WHERE name LIKE ?"
                params.append(f'%{search_term}%')
            cursor.execute(query, params)
            organizations = cursor.fetchall()
            conn.close()
            return jsonify([dict(org) for org in organizations]), 200
        except Exception as e:
            return jsonify({"message": f"Error fetching organizations: {str(e)}"}), 500

    if request.method == 'POST':
        data = request.get_json()
        company_name = data.get('company_name')
        name = data.get('name')
        if not all([company_name, name]):
            return jsonify({"message": "Company name and organization name are required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("INSERT INTO organizations (name, industry, phone, email, address, description) VALUES (?, ?, ?, ?, ?, ?)",
                           (name, data.get('industry'), data.get('phone'), data.get('email'), data.get('address'), data.get('description')))
            conn.commit()
            new_org_id = cursor.lastrowid
            conn.close()
            return jsonify({"message": "Organization added successfully", "id": new_org_id, "name": name}), 201
        except sqlite3.IntegrityError:
            return jsonify({"message": "Organization with this name already exists"}), 409
        except Exception as e:
            return jsonify({"message": f"Error adding organization: {str(e)}"}), 500

@app.route('/api/organizations/<int:org_id>', methods=['GET', 'PUT', 'DELETE'])
def handle_organization(org_id):
    """Handles GET, PUT, and DELETE for a single organization."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        if not company_name: return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM organizations WHERE id = ?", (org_id,))
            org = cursor.fetchone()
            conn.close()
            return jsonify(dict(org)) if org else (jsonify({"message": "Organization not found"}), 404)
        except Exception as e:
            return jsonify({"message": f"Error fetching organization: {str(e)}"}), 500

    elif request.method == 'PUT':
        data = request.get_json()
        company_name = data.get('company_name')
        name = data.get('name')
        if not all([company_name, name]):
            return jsonify({"message": "Company name and organization name are required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("UPDATE organizations SET name = ?, industry = ?, phone = ?, email = ?, address = ?, description = ? WHERE id = ?",
                           (name, data.get('industry'), data.get('phone'), data.get('email'), data.get('address'), data.get('description'), org_id))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "Organization not found"}), 404
            conn.close()
            return jsonify({"message": "Organization updated successfully"}), 200
        except sqlite3.IntegrityError:
            return jsonify({"message": "Organization with this name already exists"}), 409
        except Exception as e:
            return jsonify({"message": f"Error updating organization: {str(e)}"}), 500

    elif request.method == 'DELETE':
        company_name = request.args.get('company_name')
        if not company_name: return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("UPDATE contacts SET organization_id = NULL WHERE organization_id = ?", (org_id,))
            cursor.execute("DELETE FROM organizations WHERE id = ?", (org_id,))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "Organization not found"}), 404
            conn.close()
            return jsonify({"message": "Organization and related contacts updated successfully"}), 200
        except Exception as e:
            return jsonify({"message": f"Error deleting organization: {str(e)}"}), 500


# --- Contact Endpoints ---
@app.route('/api/contacts', methods=['GET', 'POST'])
def handle_contacts():
    """Handles GET and POST for contacts."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        search_term = request.args.get('search', '')
        organization_id = request.args.get('organization_id')
        if not company_name: return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            query = """
                SELECT c.*, o.name AS organization_name FROM contacts c
                LEFT JOIN organizations o ON c.organization_id = o.id
            """
            conditions, params = [], []
            if search_term:
                conditions.append("(c.first_name LIKE ? OR c.last_name LIKE ? OR o.name LIKE ?)")
                params.extend([f'%{search_term}%'] * 3)
            if organization_id:
                conditions.append("c.organization_id = ?")
                params.append(organization_id)
            if conditions:
                query += " WHERE " + " AND ".join(conditions)
            cursor.execute(query, params)
            contacts = cursor.fetchall()
            conn.close()
            return jsonify([dict(c) for c in contacts]), 200
        except Exception as e:
            return jsonify({"message": f"Error fetching contacts: {str(e)}"}), 500

    if request.method == 'POST':
        data = request.get_json()
        company_name = data.get('company_name')
        first_name = data.get('first_name')
        last_name = data.get('last_name')
        if not all([company_name, first_name, last_name]):
            return jsonify({"message": "Company name, first name, and last name are required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("INSERT INTO contacts (organization_id, first_name, last_name, title, phone, email, notes) VALUES (?, ?, ?, ?, ?, ?, ?)",
                           (data.get('organization_id'), first_name, last_name, data.get('title'), data.get('phone'), data.get('email'), data.get('notes')))
            conn.commit()
            new_contact_id = cursor.lastrowid
            conn.close()
            return jsonify({"message": "Contact added successfully", "id": new_contact_id}), 201
        except Exception as e:
            return jsonify({"message": f"Error adding contact: {str(e)}"}), 500

@app.route('/api/contacts/<int:contact_id>', methods=['GET', 'PUT', 'DELETE'])
def handle_contact(contact_id):
    """Handles GET, PUT, DELETE for a single contact."""
    if request.method == 'GET':
        company_name = request.args.get('company_name')
        if not company_name: return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM contacts WHERE id = ?", (contact_id,))
            contact = cursor.fetchone()
            conn.close()
            return jsonify(dict(contact)) if contact else (jsonify({"message": "Contact not found"}), 404)
        except Exception as e:
            return jsonify({"message": f"Error fetching contact: {str(e)}"}), 500

    elif request.method == 'PUT':
        data = request.get_json()
        company_name = data.get('company_name')
        first_name = data.get('first_name')
        last_name = data.get('last_name')
        if not all([company_name, first_name, last_name]):
            return jsonify({"message": "Company name, first name, and last name are required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("UPDATE contacts SET organization_id = ?, first_name = ?, last_name = ?, title = ?, phone = ?, email = ?, notes = ? WHERE id = ?",
                           (data.get('organization_id'), first_name, last_name, data.get('title'), data.get('phone'), data.get('email'), data.get('notes'), contact_id))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "Contact not found"}), 404
            conn.close()
            return jsonify({"message": "Contact updated successfully"}), 200
        except Exception as e:
            return jsonify({"message": f"Error updating contact: {str(e)}"}), 500

    elif request.method == 'DELETE':
        company_name = request.args.get('company_name')
        if not company_name: return jsonify({"message": "Company name is required"}), 400
        try:
            conn = get_db_connection(company_name)
            cursor = conn.cursor()
            cursor.execute("DELETE FROM contacts WHERE id = ?", (contact_id,))
            conn.commit()
            if cursor.rowcount == 0:
                conn.close()
                return jsonify({"message": "Contact not found"}), 404
            conn.close()
            return jsonify({"message": "Contact deleted successfully"}), 200
        except Exception as e:
            return jsonify({"message": f"Error deleting contact: {str(e)}"}), 500


# --- Letter Endpoints ---
@app.route('/api/letters/generate', methods=['POST'])
def generate_letter():
    """Generates a DOCX letter using a template, fills placeholders, and saves it locally."""
    data = request.get_json()
    company_name = data.get('company_name')
    subject = data.get('subject')
    body = data.get('body')
    letter_type = data.get('letter_type')
    organization_id = data.get('organization_id')
    contact_id = data.get('contact_id')
    user_id = data.get('user_id')

    if not all([company_name, subject, body, letter_type, user_id]):
        return jsonify({"message": "All fields are required"}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM company_settings WHERE company_name = ?", (company_name,))
        settings = cursor.fetchone()
        
        if not settings or not settings['letter_template_path'] or not os.path.exists(settings['letter_template_path']):
            conn.close()
            return jsonify({"message": "Letter template not configured or not found."}), 400

        current_gregorian_date = datetime.now()
        date_shamsi_persian_full = jdatetime.date.fromgregorian(date=current_gregorian_date).strftime("%Y/%m/%d")
        date_shamsi_yymmdd = jdatetime.date.fromgregorian(date=current_gregorian_date).strftime("%y%m%d")

        cursor.execute("SELECT MAX(seq_num) FROM letters WHERE company_abbr = ? AND date_shamsi_persian = ?",
                       (settings['company_short_name'], date_shamsi_persian_full))
        max_seq_num = cursor.fetchone()[0]
        next_seq_num = (max_seq_num or 0) + 1
        letter_code_persian = f"{settings['company_short_name']}-{letter_type}-{date_shamsi_yymmdd}-{next_seq_num:03d}"

        org_name_row = cursor.execute("SELECT name FROM organizations WHERE id = ?", (organization_id,)).fetchone() if organization_id else None
        contact_info_row = cursor.execute("SELECT first_name, last_name FROM contacts WHERE id = ?", (contact_id,)).fetchone() if contact_id else None
        
        doc = Document(settings['letter_template_path'])
        placeholders = {
            '[[DATE]]': date_shamsi_persian_full,
            '[[CODE]]': letter_code_persian,
            '[[ORGANIZATION_NAME]]': org_name_row['name'] if org_name_row else '',
            '[[CONTACT_NAME]]': f"{contact_info_row['first_name']} {contact_info_row['last_name']}" if contact_info_row else '',
            '[[SUBJECT]]': subject,
            '[[BODY]]': body,
            '[[COMPANY_NAME]]': settings['company_full_name_footer']
        }

        for p in doc.paragraphs:
            for key, value in placeholders.items():
                replace_placeholder_in_paragraph(p, key, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in placeholders.items():
                            replace_placeholder_in_paragraph(p, key, value)

        generated_dir = os.path.join(GENERATED_LETTERS_BASE_DIR, company_name)
        os.makedirs(generated_dir, exist_ok=True)
        local_file_path = os.path.join(generated_dir, f"{letter_code_persian}.docx")
        doc.save(local_file_path)
        
        cursor.execute("""
            INSERT INTO letters (company_abbr, seq_num, letter_code_persian, type, date_shamsi_persian, subject, body, organization_id, contact_id, local_file_path, current_gregorian_date, user_id)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            settings['company_short_name'], next_seq_num, letter_code_persian, letter_type, date_shamsi_persian_full,
            subject, body, organization_id, contact_id, local_file_path,
            current_gregorian_date.strftime("%Y-%m-%d %H:%M:%S"), user_id
        ))
        conn.commit()
        letter_id = cursor.lastrowid
        conn.close()

        final_doc = Document(local_file_path)
        full_letter_text = "\n".join([p.text for p in final_doc.paragraphs])
        
        return jsonify({
            "message": "Letter generated successfully",
            "letter_code": letter_code_persian,
            "letter_id": letter_id,
            "download_url": url_for('download_letter', letter_id=letter_id, _external=True, company_name=company_name),
            "letter_content": full_letter_text
        }), 201
    except Exception as e:
        print(f"Error in generate_letter: {e}")
        return jsonify({"message": f"An error occurred: {e}"}), 500

@app.route('/api/letters', methods=['GET'])
def get_letters():
    """Retrieves letters for a company."""
    company_name = request.args.get('company_name')
    search_term = request.args.get('search', '')
    if not company_name: return jsonify({"message": "Company name is required"}), 400
    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        query = """
            SELECT l.*, o.name AS organization_name, c.first_name, c.last_name FROM letters l
            LEFT JOIN organizations o ON l.organization_id = o.id
            LEFT JOIN contacts c ON l.contact_id = c.id
        """
        params = []
        if search_term:
            query += " WHERE l.letter_code_persian LIKE ? OR l.subject LIKE ? OR o.name LIKE ?"
            params.extend([f'%{search_term}%'] * 3)
        cursor.execute(query, params)
        letters = cursor.fetchall()
        conn.close()
        return jsonify([dict(l) for l in letters]), 200
    except Exception as e:
        return jsonify({"message": f"Error fetching letters: {str(e)}"}), 500

@app.route('/api/letters/<int:letter_id>', methods=['GET'])
def get_letter(letter_id):
    """Retrieves a single letter by ID."""
    company_name = request.args.get('company_name')
    if not company_name: return jsonify({"message": "Company name is required"}), 400
    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("""
            SELECT l.*, o.name AS organization_name, c.first_name, c.last_name FROM letters l
            LEFT JOIN organizations o ON l.organization_id = o.id
            LEFT JOIN contacts c ON l.contact_id = c.id
            WHERE l.id = ?
        """, (letter_id,))
        letter = cursor.fetchone()
        conn.close()
        
        if not letter:
            return jsonify({"message": "Letter not found"}), 404
        
        letter_content_text = ""
        if letter['local_file_path'] and os.path.exists(letter['local_file_path']):
            doc = Document(letter['local_file_path'])
            letter_content_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            letter_content_text = "فایل نامه یافت نشد."

        letter_dict = dict(letter)
        letter_dict['letter_content'] = letter_content_text
        
        return jsonify(letter_dict), 200
    except Exception as e:
        return jsonify({"message": f"Error fetching letter details: {str(e)}"}), 500

@app.route('/api/letters/download/<int:letter_id>', methods=['GET'])
def download_letter(letter_id):
    """Serves the generated DOCX letter file for download."""
    company_name = request.args.get('company_name')
    if not company_name: return jsonify({"message": "Company name is required"}), 400
    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        cursor.execute("SELECT local_file_path FROM letters WHERE id = ?", (letter_id,))
        letter = cursor.fetchone()
        conn.close()
        if not letter or not letter['local_file_path'] or not os.path.exists(letter['local_file_path']):
            return jsonify({"message": "Letter file not found"}), 404
        
        directory, filename = os.path.split(letter['local_file_path'])
        encoded_filename = urllib.parse.quote(filename.encode('utf-8'))
        response = send_from_directory(directory, filename, as_attachment=True)
        response.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{encoded_filename}"
        return response
    except Exception as e:
        return jsonify({"message": f"Error serving letter file: {str(e)}"}), 500

# --- NEW: Reporting Endpoints ---
@app.route('/api/reports/letters-by-period', methods=['GET'])
def get_letters_by_period_report():
    """Generates a report of letters created per period (daily, weekly, monthly, yearly)."""
    company_name = request.args.get('company_name')
    period = request.args.get('period', 'monthly')

    if not company_name:
        return jsonify({"message": "Company name is required"}), 400
    
    if period not in ['daily', 'weekly', 'monthly', 'yearly']:
        return jsonify({"message": "Invalid period specified"}), 400

    try:
        conn = get_db_connection(company_name)
        cursor = conn.cursor()
        
        # Use the Persian date column directly
        cursor.execute("SELECT date_shamsi_persian FROM letters")
        dates = cursor.fetchall()
        conn.close()

        counts = collections.defaultdict(int)
        for row in dates:
            try:
                # The date is already in 'YYYY/MM/DD' format
                j_date_str = row['date_shamsi_persian']
                year, month, day = map(int, j_date_str.split('/'))
                j_date = jdatetime.date(year, month, day)
                
                if period == 'daily':
                    key = j_date.strftime('%Y-%m-%d')
                elif period == 'weekly':
                    # Saturday is the start of the week in jdatetime (weekday() == 0)
                    start_of_week = j_date - jdatetime.timedelta(days=((j_date.weekday() + 1) % 7))
                    key = f"هفته {start_of_week.strftime('%Y/%m/%d')}"
                elif period == 'monthly':
                    key = j_date.strftime('%Y-%m') # e.g., '1403-04'
                elif period == 'yearly':
                    key = j_date.strftime('%Y')   # e.g., '1403'
                
                counts[key] += 1
            except (ValueError, TypeError, AttributeError):
                # Skip if the date format is incorrect or null
                continue
        
        # Sort the dictionary by keys (periods)
        sorted_counts = dict(sorted(counts.items()))

        return jsonify(sorted_counts), 200

    except Exception as e:
        print(f"Error generating report: {e}")
        return jsonify({"message": f"Error generating report: {str(e)}"}), 500
# --- Main execution block ---
if __name__ == '__main__':
    app.run(debug=True, port=5000)
